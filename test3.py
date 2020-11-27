from tkinter import *
import cx_Oracle
import itertools
from docx import Document
from docx.shared import Inches
from tkinter import messagebox
from docx.enum.section import WD_ORIENT
import os

connstr = 'fbla_csa/fbla_csa@fbladb1'

labels = []
buttons = []
tboxes = []
dropdowns = []

# Used to refresh the page

def destroy_all():
   for label in labels:
      label.destroy()

   for button in buttons:
      button.destroy()

   for tbox in tboxes:
      tbox.destroy()

   for dropdown in dropdowns:
      dropdown.destroy()
      
#def clear_frame(frame):
#   frame.grid_forget()

# Does nothing

def donothing():
   filewin = Toplevel(root)
   button = Button(filewin, text="Do nothing button")
   button.pack()

# Queries student data in the query forms

def query_student():
   destroy_all()
   #clear_button1 = Button(root, text='Clear Workspace',command=destroy_all)
   #clear_button1.grid(row=0,column=10)   
   #buttons.append(clear_button1)
   conn1 = cx_Oracle.connect(connstr)
   cur1 = conn1.cursor()
   cur1.execute("select * from student")
   result1 = cur1.fetchall()

   column1 = Label(root,text='Student ID')
   column1.grid(row=0,column=0)
   labels.append(column1)
   column2 = Label(root,text='First Name')
   column2.grid(row=0,column=1)
   labels.append(column2)
   column3 = Label(root,text='Middle Name')
   column3.grid(row=0,column=2)
   labels.append(column3)
   column4 = Label(root,text='Last Name')
   column4.grid(row=0,column=3)
   labels.append(column4)
   column5 = Label(root,text='Grade')
   column5.grid(row=0,column=4)
   labels.append(column5)
   column6 = Label(root,text='School Name')
   column6.grid(row=0,column=5)
   labels.append(column6)
   column7 = Label(root,text='Chapter ID')
   column7.grid(row=0,column=6)
   labels.append(column7)
   column8 = Label(root,text='Gender')
   column8.grid(row=0,column=7)
   labels.append(column8)
   column9 = Label(root,text='Date of Birth')
   column9.grid(row=0,column=8)
   labels.append(column9)
   
   for index1, x1 in enumerate(result1):
      num1 = 0
      index1 += 2
      for y1 in x1: 
         lookup_label = Label(root, text=y1)
         labels.append(lookup_label)
         lookup_label.grid(row=index1,column=num1)
         num1 += 1

# Queries FBLA chapter info in the query forms

def query_chapter():
   destroy_all()
   #clear_button2 = Button(root, text='Clear Workspace',command=destroy_all)
   #clear_button2.grid(row=0,column=10)
   #buttons.append(clear_button2)
   conn2 = cx_Oracle.connect(connstr)
   cur2 = conn2.cursor()
   cur2.execute('select * from chapter')
   result2 = cur2.fetchall()
   
   column1 = Label(root,text='Chapter ID')
   column1.grid(row=0,column=0)
   labels.append(column1)
   column2 = Label(root,text='Chapter Name')
   column2.grid(row=0,column=1)
   labels.append(column2)
   column3 = Label(root,text='Address')
   column3.grid(row=0,column=2)
   labels.append(column3)
   column4 = Label(root,text='City')
   column4.grid(row=0,column=3)
   labels.append(column4)
   column5 = Label(root,text='Zipcode')
   column5.grid(row=0,column=4)
   labels.append(column5)
   column6 = Label(root,text='County')
   column6.grid(row=0,column=5)
   labels.append(column6)
   column7 = Label(root,text='Country')
   column7.grid(row=0,column=6)
   labels.append(column7)
   
   for index2, x2 in enumerate(result2):
      num2 = 0
      index2 += 2
      for y2 in x2:
         chap_label = Label(root, text=y2)
         labels.append(chap_label)
         chap_label.grid(row=index2,column=num2)
         num2 += 1
#         clear_button = Button(root, text='Clear Workspace', command=lambda: clear_button.grid_remove())
#         clear_button.place(x=900,y=200)

# Queries program categories in the query forms

def query_progcat():
   destroy_all()
   #clear_button3 = Button(root, text='Clear Workspace',command=destroy_all)
   #clear_button3.grid(row=0,column=10)
   #buttons.append(clear_button3)
   conn3 = cx_Oracle.connect(connstr)
   cur3 = conn3.cursor()
   cur3.execute('select * from program_category')
   result3 = cur3.fetchall()

   column1 = Label(root,text='Program Category ID')
   column1.grid(row=0,column=0)
   labels.append(column1)
   column2 = Label(root,text='Name')
   column2.grid(row=0,column=1)
   labels.append(column2)
   column3 = Label(root,text='Minimum Hours')
   column3.grid(row=0,column=2)
   labels.append(column3)
   
   for index3, x3 in enumerate(result3):
      num3 = 0
      index3 += 2
      for y3 in x3:
         prog_label = Label(root, text=y3)
         labels.append(prog_label)
         prog_label.grid(row=index3,column=num3)
         num3 += 1

# Queries community service awards in the query forms

def query_csa():
   destroy_all()
   #clear_button4 = Button(root, text='Clear Workspace',command=destroy_all)
   #clear_button4.grid(row=0,column=15)
   #buttons.append(clear_button4)
   conn4 = cx_Oracle.connect(connstr)
   cur4 = conn4.cursor()
   cur4.execute('select * from community_service')
   result4 = cur4.fetchall()

   column1 = Label(root,text='Community Service ID')
   column1.grid(row=0,column=0)
   labels.append(column1)
   column2 = Label(root,text='Name')
   column2.grid(row=0,column=1)
   labels.append(column2)
   column3 = Label(root,text='Address')
   column3.grid(row=0,column=2)
   labels.append(column3)
   column4 = Label(root,text='Contact')
   column4.grid(row=0,column=3)
   labels.append(column4)
   column5 = Label(root,text='Address')
   column5.grid(row=0,column=4)
   labels.append(column5)
   column6 = Label(root,text='City')
   column6.grid(row=0,column=5)
   labels.append(column6)
   column7 = Label(root,text='Zipcode')
   column7.grid(row=0,column=6)
   labels.append(column7)
   column8 = Label(root,text='County')
   column8.grid(row=0,column=7)
   labels.append(column8)
   column9 = Label(root,text='Country')
   column9.grid(row=0,column=8)
   labels.append(column9)
   column10 = Label(root,text='Work Phone')
   column10.grid(row=0,column=9)
   labels.append(column10)
   column11 = Label(root,text='Cell Phone')
   column11.grid(row=0,column=10)
   labels.append(column11)
   column12 = Label(root,text='Email')
   column12.grid(row=0,column=11)
   labels.append(column12)
   
   for index4, x4 in enumerate(result4):
      num4 = 0
      index4 += 2
      for y4 in x4:
         csa_label = Label(root, text=y4)
         labels.append(csa_label)
         csa_label.grid(row=index4,column=num4)
         num4 += 1

# Queries community service hours in the query forms

def query_csa_hour():
   destroy_all()
   #clear_button5 = Button(root, text='Clear Workspace',command=destroy_all)
   #clear_button5.grid(row=0,column=8)
   #buttons.append(clear_button5)
   conn5 = cx_Oracle.connect(connstr)
   cur5 = conn5.cursor()
   cur5.execute('select * from csa_hour')
   result5 = cur5.fetchall()
   
   column1 = Label(root,text='Community Service Hour ID')
   column1.grid(row=0,column=0)
   labels.append(column1)
   column2 = Label(root,text='Student ID')
   column2.grid(row=0,column=1)
   labels.append(column2)
   column3 = Label(root,text='Program Category ID')
   column3.grid(row=0,column=2)
   labels.append(column3)
   column4 = Label(root,text='Community Service ID')
   column4.grid(row=0,column=3)
   labels.append(column4)
   column5 = Label(root,text='Hours Worked')
   column5.grid(row=0,column=4)
   labels.append(column5)
   column6 = Label(root,text='Date Worked')
   column6.grid(row=0,column=5)
   labels.append(column6)
   
   for index5, x5 in enumerate(result5):
      num5 = 0
      index5 += 2
      for y5 in x5:
         hour_label = Label(root, text=y5)
         labels.append(hour_label)
         hour_label.grid(row=index5,column=num5)
         num5 += 1

# Queries award categories in the query forms

def query_award_det():
   destroy_all()
   conn6 = cx_Oracle.connect(connstr)
   cur6 = conn6.cursor()
   cur6.execute('select * from award_detail')
   result6 = cur6.fetchall()
   print(enumerate(result6))

   column1 = Label(root,text='Program Category ID')
   column1.grid(row=0,column=0)
   labels.append(column1)
   column2 = Label(root,text='Student ID')
   column2.grid(row=0,column=1)
   labels.append(column2)
   column3 = Label(root,text='Community Service ID')
   column3.grid(row=0,column=2)
   labels.append(column3)
   column4 = Label(root,text='Has Award?')
   column4.grid(row=0,column=3)
   labels.append(column4)

   for index6, x6 in enumerate(result6):
      num6 = 0
      index6 += 2
      for y6 in x6:
         query_label = Label(root, text=y6)
         labels.append(query_label)
         query_label.grid(row=index6,column=num6)
         num6 += 1

# Adds New Student in Add Entry Forms

def add_student():
   destroy_all()

   def clear():
      sid_entry.delete(0,END)
      fname_entry.delete(0,END)
      mname_entry.delete(0,END)
      lname_entry.delete(0,END)
      school_entry.delete(0,END)
      chap_id_entry.delete(0,END)
      dob_entry.delete(0,END)
      
   def register_student():
      conn7 = cx_Oracle.connect(connstr)
      cur7 = conn7.cursor()
      command1 = f"""
               insert into student values ({sid.get()},'{fname.get()}',
               '{mname.get()}','{lname.get()}',
               '{grade.get()}','{school_name.get()}','{chap_id.get()}',
               '{gender.get()}','{dob.get()}')"""
      print(command1)
      cur7.execute(command1)
      conn7.commit()
      cur7.close()

   # Variables
   global gender
   sid = IntVar()
   fname = StringVar()
   mname = StringVar()
   lname = StringVar()
   grade = IntVar()
   school_name = StringVar()
   chap_id = IntVar()
   gender = StringVar()
   dob = StringVar()

   GRADES = [9,10,11,12]
   GENDER = ['M','F']
   
   # Label Variables
   title = Label(root,text='Add New Student',font=('Times New Roman',15))
   labels.append(title)
   title.place(x=180,y=15)
   
   sid_text = Label(root, text='Student ID:')
   labels.append(sid_text)
   sid_text.place(x=100,y=50)
   
   fname_text = Label(root, text='First Name:')
   labels.append(fname_text)
   fname_text.place(x=100,y=100)
   
   mname_text = Label(root,text='Middle Name:')
   labels.append(mname_text)
   mname_text.place(x=100,y=150)
   
   lname_text = Label(root,text='Last Name:')
   labels.append(lname_text)
   lname_text.place(x=100,y=200)
   
   grade_text = Label(root,text='Grade:')
   labels.append(grade_text)
   grade_text.place(x=100,y=250)

   school_name_text = Label(root,text='School Name:')
   labels.append(school_name_text)
   school_name_text.place(x=100,y=300)

   chap_id_text = Label(root,text='Chapter ID:')
   labels.append(chap_id_text)
   chap_id_text.place(x=100,y=350)

   gender_text = Label(root,text='Gender:')
   labels.append(gender_text)
   gender_text.place(x=100,y=400)

   dob_text = Label(root,text='Date of Birth (Example ~ 01-JAN-2000):')
   labels.append(dob_text)
   dob_text.place(x=100,y=450)

   # Entry Variables
   sid_entry = Entry(textvariable=sid,width=20)
   tboxes.append(sid_entry)
   sid_entry.place(x=200,y=50)

   fname_entry = Entry(textvariable=fname,width=20)
   tboxes.append(fname_entry)
   fname_entry.place(x=200,y=100)

   mname_entry = Entry(textvariable=mname,width=20)
   tboxes.append(mname_entry)
   mname_entry.place(x=200,y=150)

   lname_entry = Entry(textvariable=lname,width=20)
   tboxes.append(lname_entry)
   lname_entry.place(x=200,y=200)

   grade_entry = OptionMenu(root,grade,*GRADES)
   tboxes.append(grade_entry)
   grade_entry.place(x=200,y=250)

   school_entry = Entry(textvariable=school_name,width=20)
   tboxes.append(school_entry)
   school_entry.place(x=200,y=300)

   chap_id_entry = Entry(textvariable=chap_id,width=20)
   tboxes.append(chap_id_entry)
   chap_id_entry.place(x=200,y=350)

   gender_entry = OptionMenu(root,gender,*GENDER)
   tboxes.append(gender_entry)
   gender_entry.place(x=200,y=400)

   dob_entry = Entry(textvariable=dob,width=20)
   tboxes.append(dob_entry)
   dob_entry.place(x=350,y=450)

   # Register Button
   register1 = Button(root,text='Register',command=register_student)
   tboxes.append(register1)
   register1.place(x=150,y=500)

   # Clear Button
   clear1 = Button(root,text='Clear All Fields',command=clear)
   tboxes.append(clear1)
   clear1.place(x=250,y=500)

def add_chapter():
   destroy_all()

   conn1 = cx_Oracle.connect(connstr)
   curs1 = conn1.cursor()
   curs1.execute('select chapter_id from chapter')
   CHAPTER_ID_IN = curs1.fetchall()
   CHAPTER_ID = list(itertools.chain(*CHAPTER_ID_IN))
   print(CHAPTER_ID)

   curs1.execute('select name from chapter')
   SCHOOL_NAMES = curs1.fetchall()
   print(SCHOOL_NAMES)

   def register():
      curs1.execute(f'''insert into chapter values ('{cid.get()}',
                     '{name.get()}','{address.get()}','{city.get()}',
                     '{zipcode.get()}','{county.get()}','{country.get()}')''')

      conn1.commit()

   def clear():
      cid_entry.delete(0,END)
      name_entry.delete(0,END)
      address_entry.delete(0,END)
      city_entry.delete(0,END)
      zipcode_entry.delete(0,END)
      county_entry.delete(0,END)
      country_entry.delete(0,END)
      
   # Variables
   cid = IntVar()
   name = StringVar()
   address = StringVar()
   city = StringVar()
   zipcode = IntVar()
   county = StringVar()
   country = StringVar()

   # Label Variables
   title = Label(root,text='Add New Chapter',font=('Times New Roman',15))
   labels.append(title)
   title.place(x=180,y=15)
   
   cid_lab = Label(root,text='Chapter ID:')
   labels.append(cid_lab)
   name_lab = Label(root,text='Chapter Name:')
   labels.append(name_lab)
   address_lab = Label(root,text='Address:')
   labels.append(address_lab)
   city_lab = Label(root,text='City Name:')
   labels.append(city_lab)
   zipcode_lab = Label(root,text='Zipcode:')
   labels.append(zipcode_lab)
   county_lab = Label(root,text='County:')
   labels.append(county_lab)
   country_lab = Label(root,text='Country:')
   labels.append(country_lab)

   # Place label variables
   cid_lab.place(x=100,y=50)
   name_lab.place(x=100,y=100)
   address_lab.place(x=100,y=150)
   city_lab.place(x=100,y=200)
   zipcode_lab.place(x=100,y=250)
   county_lab.place(x=100,y=300)
   country_lab.place(x=100,y=350)

   # Entry Variables
   cid_entry = Entry(textvariable=cid,width=20)
   tboxes.append(cid_entry)
   name_entry = Entry(textvariable=name,width=20)
   tboxes.append(name_entry)
   address_entry = Entry(textvariable=address,width=20)
   tboxes.append(address_entry)
   city_entry = Entry(textvariable=city,width=20)
   tboxes.append(city_entry)
   zipcode_entry = Entry(textvariable=zipcode,width=20)
   tboxes.append(zipcode_entry)
   county_entry = Entry(textvariable=county,width=20)
   tboxes.append(county_entry)
   country_entry = Entry(textvariable=country,width=20)
   tboxes.append(country_entry)

   # Place entry variables
   cid_entry.place(x=200,y=50)
   name_entry.place(x=200,y=100)
   address_entry.place(x=200,y=150)
   city_entry.place(x=200,y=200)
   zipcode_entry.place(x=200,y=250)
   county_entry.place(x=200,y=300)
   country_entry.place(x=200,y=350)

   if cid in CHAPTER_ID:
      messagebox.showerror('Chapter ID exists!','Choose a different Chapter ID!')
      

   # Create and place the register button
   register2 = Button(root,text='Add Chapter',command=register)
   tboxes.append(register2)
   register2.place(x=100,y=400)

   # Create and place the clear button
   clear1 = Button(root,text='Clear All Fields',command=clear)
   tboxes.append(clear1)
   clear1.place(x=200,y=400)

def add_award():
   destroy_all()
   
   conn2 = cx_Oracle.connect(connstr)
   curs2 = conn2.cursor()
   curs2.execute('select program_category_id from program_category')
   CAT_ID_IN = curs2.fetchall()
   print(CAT_ID_IN)
   curs2.execute('select student_id from student')
   STUDENT_ID_IN = curs2.fetchall()
   curs2.execute('select community_service_id from community_service')
   CS_ID_IN = curs2.fetchall()
   HAS_AWARD = ['YES','NO']

   CAT_ID = [item for t in CAT_ID_IN for item in t]
   print(CAT_ID_IN)
   STUDENT_ID = list(itertools.chain(*STUDENT_ID_IN))
   CS_ID = list(itertools.chain(*CS_ID_IN))

   def clear():
   #   cat_id_entry['menu'].delete(0,END)
   #   student_id_entry['menu'].delete(0,END)
   #   cs_id_entry['menu'].delete(0,END)
   #   has_award_entry['menu'].delete(0,END)
      pass

    
   def register_award():
      command = f"insert into award_detail values ('{cat_id.get()}','{student_id.get()}','{cs_id.get()}','Y')"
      print(command)
      curs2.execute(command)
      conn2.commit()
   
   # Variables
   cat_id = StringVar()
   student_id = StringVar()
   cs_id = StringVar()
   has_award = StringVar()

   # Label Variables
   title = Label(root,text='Add New Award Details',font=('Times New Roman',15))
   labels.append(title)
   title.place(x=180,y=15)
   
   cat_id_label = Label(root,text='Program Category:')
   labels.append(cat_id_label)
   student_id_label = Label(root,text='Student ID:')
   labels.append(student_id_label)
   cs_id_label = Label(root,text='Community Service ID:')
   labels.append(cs_id_label)
   has_award_label = Label(root,text='Does the student have an award already?')
   labels.append(has_award_label)
   
   # Place lavel variables
   cat_id_label.place(x=100,y=50)
   student_id_label.place(x=100,y=100)
   cs_id_label.place(x=100,y=150)
   has_award_label.place(x=100,y=200)
   
   # Entry Variables
   cat_id_entry = OptionMenu(root,cat_id,*CAT_ID)
   labels.append(cat_id_entry)
   student_id_entry = OptionMenu(root,student_id,*STUDENT_ID)
   labels.append(student_id_entry)
   cs_id_entry = OptionMenu(root,cs_id,*CS_ID)
   labels.append(cs_id_entry)
   has_award_entry = OptionMenu(root,has_award,*HAS_AWARD)
   labels.append(has_award_entry)
   
   # Place entry variables
   cat_id_entry.place(x=250,y=50)
   student_id_entry.place(x=250,y=100)
   cs_id_entry.place(x=250,y=150)
   has_award_entry.place(x=350,y=200)

   # Create and place a register button
   register3 = Button(root,text='Add Award Details',command=register_award)
   labels.append(register3)
   register3.place(x=100,y=250)
   
   # Create and place a button to clear all fields
   clear = Button(root,text='Clear All Fields',command=clear)
   labels.append(clear)
   clear.place(x=250,y=250)
   
def add_cs():
   destroy_all()

   # In the insert statement, include cs_addr.get() twice!!!!
   def register():
      conn3 = cx_Oracle.connect(connstr)
      curs3 = conn3.cursor()
      command = f"""insert into community_service values ('{cs_id.get()}',
      '{cs_name.get()}','{cs_addr.get()}','{cs_contact.get()}','{cs_addr.get()}',
      '{city.get()}','{zipcode.get()}','{county.get()}','{country.get()}',
      '{wphone.get()}','{cphone.get()}','{email.get()}')"""
      curs3.execute(command)
      conn3.commit()
      curs3.close()

   def clear():
      cs_id_entry.delete(0,END)
      cs_name_entry.delete(0,END)
      cs_addr_entry.delete(0,END)
      cs_contact_entry.delete(0,END)
      city_entry.delete(0,END)
      zipcode_entry.delete(0,END)
      county_entry.delete(0,END)
      country_entry.delete(0,END)
      wphone_entry.delete(0,END)
      cphone_entry.delete(0,END)
      email_entry.delete(0,END)
      
   # Variables
   cs_id = StringVar()
   cs_name = StringVar()
   cs_addr = StringVar()
   cs_contact = StringVar()
   city = StringVar()
   zipcode = IntVar()
   county = StringVar()
   country = StringVar()
   wphone = IntVar()
   cphone = IntVar()
   email = StringVar()

   # Label Variables
   title = Label(root,text='Add New Community Service Info',font=('Times New Roman',15))
   labels.append(title)
   title.place(x=180,y=15)
   
   cs_id_label = Label(root,text='Community Service ID:')
   labels.append(cs_id_label)
   cs_name_label = Label(root,text='Name of Community Service Activity:')
   labels.append(cs_name_label)
   cs_addr_label = Label(root,text='Address of Community Service Activity:')
   labels.append(cs_addr_label)
   cs_contact_label = Label(root,text='Contact of Community Service Activity (Full Name):')
   labels.append(cs_contact_label)
   city_label = Label(root,text='City of Community Service Activity:')
   labels.append(city_label)
   zipcode_label = Label(root,text='Zipcode of Community Service Activity:')
   labels.append(zipcode_label)
   county_label = Label(root,text='County of Community Service Activity:')
   labels.append(county_label)
   country_label = Label(root,text='Country of Community Service Activity:')
   labels.append(country_label)
   wphone_label = Label(root,text='Work phone:')
   labels.append(wphone_label)
   cphone_label = Label(root,text='Cell phone:')
   labels.append(cphone_label)
   email_label = Label(root,text='Email:')
   labels.append(email_label)

   # Place labels
   cs_id_label.place(x=100,y=50)
   cs_name_label.place(x=100,y=100)
   cs_addr_label.place(x=100,y=150)
   cs_contact_label.place(x=100,y=200)
   city_label.place(x=100,y=250)
   zipcode_label.place(x=100,y=300)
   county_label.place(x=100,y=350)
   country_label.place(x=100,y=400)
   wphone_label.place(x=100,y=450)
   cphone_label.place(x=100,y=500)
   email_label.place(x=100,y=550)

   # Entry Variables
   cs_id_entry = Entry(textvariable=cs_id)
   labels.append(cs_id_entry)
   cs_name_entry = Entry(textvariable=cs_name)
   labels.append(cs_name_entry)
   cs_addr_entry = Entry(textvariable=cs_addr)
   labels.append(cs_addr_entry)
   cs_contact_entry = Entry(textvariable=cs_contact)
   labels.append(cs_contact_entry)
   city_entry = Entry(textvariable=city)
   labels.append(city_entry)
   zipcode_entry = Entry(textvariable=zipcode)
   labels.append(zipcode_entry)
   county_entry = Entry(textvariable=county)
   labels.append(county_entry)
   country_entry = Entry(textvariable=country)
   labels.append(country_entry)
   wphone_entry = Entry(textvariable=wphone)
   labels.append(wphone_entry)
   cphone_entry = Entry(textvariable=cphone)
   labels.append(cphone_entry)
   email_entry = Entry(textvariable=email)
   labels.append(email_entry)
   
   # Place entry variables
   cs_id_entry.place(x=250,y=50)
   cs_name_entry.place(x=320,y=100)
   cs_addr_entry.place(x=350,y=150)
   cs_contact_entry.place(x=400,y=200)
   city_entry.place(x=325,y=250)
   zipcode_entry.place(x=325,y=300)
   county_entry.place(x=325,y=350)
   country_entry.place(x=325,y=400)
   wphone_entry.place(x=200,y=450)
   cphone_entry.place(x=200,y=500)
   email_entry.place(x=180,y=550)
   
   # Create and place register button
   register = Button(root,text='Add Community Service Activity Info',command=register)
   labels.append(register)
   register.place(x=100,y=600)

   # Clear button
   clear = Button(root,text='Clear All Fields',command=clear)
   labels.append(clear)
   clear.place(x=350,y=600)
   
def add_cshrs():
   destroy_all()

   conn4 = cx_Oracle.connect(connstr)
   curs4 = conn4.cursor()
   curs4.execute('select student_id from student')
   STUDENT_ID_IN = curs4.fetchall()
   curs4.execute('select program_category_id from program_category')
   PROGCAT_ID_IN = curs4.fetchall()
   curs4.execute('select community_service_id from community_service')
   CS_ID_IN = curs4.fetchall()

   STUDENT_ID = [item for t in STUDENT_ID_IN for item in t]
   PROGCAT_ID = [item for u in PROGCAT_ID_IN for item in u]
   CS_ID = [item for v in CS_ID_IN for item in v]

   def register():
      command = f"""insert into csa_hour values ('{hr_id.get()}',
               '{student_id.get()}','{progcat_id.get()}','{cs_id.get()}',
               '{hours.get()}','{date.get()}')"""
      print(command)
      curs4.execute(command)
      conn4.commit()
      
   def clear():
      hr_id_entry.delete(0,END)
      hours_entry.delete(0,END)
      date_entry.delete(0,END)

   # Variables
   hr_id = StringVar()
   student_id = StringVar()
   progcat_id = StringVar()
   cs_id = StringVar()
   hours = StringVar()
   date = StringVar()

   # Label Variables
   title = Label(root,text='Add Community Service Hours',font=('Times New Roman',15))
   labels.append(title)
   title.place(x=180,y=15)
   
   label1 = Label(root,text='Hour ID:')
   labels.append(label1)
   label2 = Label(root,text='Student ID:')
   labels.append(label2)
   label3 = Label(root,text='Program Category ID:')
   labels.append(label3)
   label4 = Label(root,text='Community Service ID:')
   labels.append(label4)
   label5 = Label(root,text='Hours Worked:')
   labels.append(label5)
   label6 = Label(root,text='Date of Service:')
   labels.append(label6)
   
   # Place Label Variables
   label1.place(x=100,y=50)
   label2.place(x=100,y=100)
   label3.place(x=100,y=150)
   label4.place(x=100,y=200)
   label5.place(x=100,y=250)
   label6.place(x=100,y=300)

   # Create entry variables
   hr_id_entry = Entry(textvariable=hr_id,width=20)
   labels.append(hr_id_entry)
   student_id_entry = OptionMenu(root,student_id,*STUDENT_ID) 
   labels.append(student_id_entry)
   progcat_id_entry = OptionMenu(root,progcat_id,*PROGCAT_ID)
   labels.append(progcat_id_entry)
   cs_id_entry = OptionMenu(root,cs_id,*CS_ID)
   labels.append(cs_id_entry)
   hours_entry = Entry(textvariable=hours,width=20)
   labels.append(hours_entry)
   date_entry = Entry(textvariable=date,width=20)
   labels.append(date_entry)

   # Place entry variables
   hr_id_entry.place(x=200,y=50)
   student_id_entry.place(x=200,y=100)
   progcat_id_entry.place(x=250,y=150)
   cs_id_entry.place(x=250,y=200)
   hours_entry.place(x=200,y=250)
   date_entry.place(x=200,y=300)

   # Register button
   register = Button(root,text='Add Community Service Hours',command=register)
   labels.append(register)
   register.place(x=100,y=350)

   # Clear button
   clear = Button(root,text='Clear All Fields',command=clear)
   labels.append(clear)
   clear.place(x=300,y=350)

#def add_progcat():
#   destroy_all()

   # Variables
   
   # Label variables
   # Place label variables
   # Entry variables
   # Place entry variables
   # Create and place Register Button

def modify_student():
   destroy_all()
   
   def search_student():

      SEARCH_MENU = ['Last Name','Student ID','Chapter ID','Grade']
      
      def search_now():
         global command
         
         searched = search_entry.get()
         if search_by.get() == 'Last Name':
            command = f"select * from student where last_name = '{searched}'"
         elif search_by.get() == 'Student ID':
            command = f"select * from student where student_id = '{searched}'"
         elif search_by.get() == 'Chapter ID':
            command = f"select * from student where chapter_id = '{searched}'"
         elif search_by.get() == 'Grade':
            command = f"select * from student where grade = '{searched}'"

         widgets = []

         def clear_labels():
            searched_label.destroy()
            for widget in widgets:
               widget.destroy()

         print(command)
         conn = cx_Oracle.connect(connstr)
         curs = conn.cursor()
         curs.execute(command)
         result = curs.fetchall()
         searched_label = Label(root,text='')
         print(result)

         for row in result:
            sid3 = row[0]
            print(f'Student ID: {sid3}')
            fname3 = row[1]
            mname3 = row[2]
            lname3 = row[3]
            grade3 = row[4]
            school_name3 = row[5]
            chap_id3 = row[6]
            gender3 = row[7]
            dob3 = row[8]

         def delete():
            command3 = f"delete from student where student_id = {sid2.get()}"
            curs.execute(command3)
            conn.commit()

         def update():
            command2 = f"""update student set first_name='{fname2.get()}',
            middle_name='{mname2.get()}',last_name='{lname2.get()}',
            grade='{grade2.get()}',school_name='{school_name2.get()}',
            chapter_id='{chap_id2.get()}',gender='{gender2.get()}',dob= TO_DATE('{dob2.get()}', 'YYYY-MM-DD HH24:MI:SS')
            where student_id='{sid2.get()}'"""

            print(command2)
            curs.execute(command2)
            conn.commit()
         
         if not result:
            searched_label.destroy()
            result = 'Record Not Found'

         else:
            widgets = []
            
            sid_text = Label(root, text='Student:')
            labels.append(sid_text)
            widgets.append(sid_text)
            sid_text.place(x=100,y=130)
            
            fname_text = Label(root, text='First Name:')
            labels.append(fname_text)
            widgets.append(fname_text)
            fname_text.place(x=100,y=180)
            
            mname_text = Label(root,text='Middle Name:')
            labels.append(mname_text)
            widgets.append(mname_text)
            mname_text.place(x=100,y=230)
            
            lname_text = Label(root,text='Last Name:')
            labels.append(lname_text)
            widgets.append(lname_text)
            lname_text.place(x=100,y=280)
            
            grade_text = Label(root,text='Grade:')
            labels.append(grade_text)
            widgets.append(grade_text)
            grade_text.place(x=100,y=330)

            school_name_text = Label(root,text='School Name:')
            labels.append(school_name_text)
            widgets.append(school_name_text)
            school_name_text.place(x=100,y=380)

            chap_id_text = Label(root,text='Chapter ID:')
            labels.append(chap_id_text)
            widgets.append(chap_id_text)
            chap_id_text.place(x=100,y=430)

            gender_text = Label(root,text='Gender:')
            labels.append(gender_text)
            widgets.append(gender_text)
            gender_text.place(x=100,y=480)

            dob_text = Label(root,text='Date of Birth (Example ~ 01-JAN-2000):')
            labels.append(dob_text)
            widgets.append(dob_text)
            dob_text.place(x=100,y=530)

            # Entry Variables
            global gender2
            sid2 = StringVar()
            fname2 = StringVar()
            mname2 = StringVar()
            lname2 = StringVar()
            grade2 = StringVar()
            school_name2 = StringVar()
            chap_id2 = StringVar()
            gender2 = StringVar()
            dob2 = StringVar()
            
            GRADES = [9,10,11,12]
            GENDER = ['M','F']
            
            sid_entry = Entry(textvariable=sid2,width=20)
            tboxes.append(sid_entry)
            widgets.append(sid_entry)
            string1 = int(result[0][0])
            sid_entry.insert(0,str(sid3))
            print(result[0][0])
            sid_entry.place(x=200,y=130)

            fname_entry = Entry(textvariable=fname2,width=20)
            tboxes.append(fname_entry)
            widgets.append(fname_entry)
            fname_entry.insert(0,fname3)
            print(result[0][1])
            fname_entry.place(x=200,y=180)

            mname_entry = Entry(textvariable=mname2,width=20)
            tboxes.append(mname_entry)
            widgets.append(mname_entry)
            #mname_entry.insert(0,result[0][2])
            print(result[0][2])
            if result[0][2] == None:
               mname_entry.insert(0,'')
            else:
               mname_entry.insert(0,mname3)
            mname_entry.place(x=200,y=230)

            lname_entry = Entry(textvariable=lname2,width=20)
            tboxes.append(lname_entry)
            widgets.append(lname_entry)
            if result[0][3] == None:
               lname_entry.insert(0,'')
            else:
               lname_entry.insert(0,lname3)
            lname_entry.place(x=200,y=280)

            grade_entry = OptionMenu(root,grade2,*GRADES)
            tboxes.append(grade_entry)
            widgets.append(grade_entry)
            grade_entry.place(x=200,y=330)

            school_entry = Entry(textvariable=school_name2,width=20)
            tboxes.append(school_entry)
            widgets.append(school_entry)
            if result[0][2] == None:
               school_entry.insert(0,'')
            else:
               school_entry.insert(0,school_name3)
            school_entry.place(x=200,y=380)

            chap_id_entry = Entry(textvariable=chap_id2,width=20)
            tboxes.append(chap_id_entry)
            widgets.append(chap_id_entry)
            if result[0][2] == None:
               chap_id_entry.insert(0,'')
            else:
               chap_id_entry.insert(0,str(chap_id3))

            chap_id_entry.place(x=200,y=430)

            gender_entry = OptionMenu(root,gender2,*GENDER)
            tboxes.append(gender_entry)
            widgets.append(gender_entry)
            gender_entry.place(x=200,y=480)

            dob_entry = Entry(textvariable=dob2,width=20)
            tboxes.append(dob_entry)
            widgets.append(dob_entry)
            if result[0][2] == None:
               dob_entry.insert(0,'')
            else:
               dob_entry.insert(0,dob3)
            dob_entry.place(x=350,y=530)

            save_record = Button(root,text='Update',command=update)
            widgets.append(save_record)
            tboxes.append(save_record)
            save_record.place(x=200,y=580)
            
            delete_button = Button(root,text='Delete Entry',command=delete)
            delete_button.place(x=300,y=580)
            tboxes.append(delete_button)
            widgets.append(delete_button)

         searched_label = Label(root,text=result)
         searched_label.grid(row=6,column=1,padx=10,pady=10,columnspan=2)
         labels.append(searched_label)
               
         clear_button = Button(root,text='Clear Search Results',command=clear_labels)
         clear_button.grid(row=0,column=4,padx=10,pady=10)
         labels.append(clear_button)
         
      def clear_labels():
         searched_label.destroy()
         
      search_entry = StringVar()
      search_by = StringVar()
      
      search_box = Entry(root,textvariable=search_entry)
      search_box.grid(row=0,column=2,padx=10,pady=10)
      labels.append(search_box)

      search_type_label = Label(root,text='Search by:')
      search_type_label.grid(row=0,column=0,padx=10,pady=10)
      
      search_box_label = OptionMenu(root,search_by,*SEARCH_MENU)
      search_box_label.grid(row=0,column=1,padx=10,pady=10)
      labels.append(search_box_label)

      search_button = Button(root,text='Search Students',command=search_now)
      search_button.grid(row=0,column=3,padx=10,pady=10)
      labels.append(search_button)

      
   search_student()


def modify_chapter():
   destroy_all()

   SEARCH_MENU = ['Chapter ID','Name']
   
   def search_chapter():
      def search_now():

         command = ''
         
         if search_by.get() == 'Chapter ID':
            command = f'select * from chapter where chapter_id={search_entry.get()}'
            print(command)
         elif search_by.get() == 'Name':
            command = f'select * from chapter where name={search_entry.get()}'
         
         print(command)
         conn = cx_Oracle.connect(connstr)
         curs = conn.cursor()
         curs.execute(command)
         result = curs.fetchall()
         searched_label = Label(root,text='')
         print(result)
         widgets = []
         
         def update():
            command2 = f"""
            update chapter set name='{name.get()}',address='{addr.get()}',city='{city.get()}',zipcode='{zipcode.get()}',
            county='{county.get()}',country='{country.get()}' where chapter_id='{cid.get()}'
            """
            print(command2)
            curs.execute(command2)
            conn.commit()

         def delete():
            command3 = f"""
            delete from chapter where chapter_id='{cid.get()}'
            """
            curs.execute(command3)
            connection.commit()

         def clear_labels():
            for widget in widgets:
               widget.destroy()

         if not result:
            result = 'No Results Found'
         else:
         
            label1 = Label(root, text='Chapter_ID:')
            labels.append(label1)
            widgets.append(label1)
            label1.place(x=100,y=130)
            
            label2 = Label(root, text='Name:')
            labels.append(label2)
            widgets.append(label2)
            label2.place(x=100,y=180)
            
            label3 = Label(root,text='Address:')
            labels.append(label3)
            widgets.append(label3)
            label3.place(x=100,y=230)
            
            label4 = Label(root,text='City:')
            labels.append(label4)
            widgets.append(label4)
            label4.place(x=100,y=280)
            
            label5 = Label(root,text='Zipcode:')
            labels.append(label5)
            widgets.append(label5)
            label5.place(x=100,y=330)

            label6 = Label(root,text='County:')
            labels.append(label6)
            widgets.append(label6)
            label6.place(x=100,y=380)

            label7 = Label(root,text='Country:')
            labels.append(label7)
            widgets.append(label7)
            label7.place(x=100,y=430)

            # Variables
            cid = StringVar()
            name = StringVar()
            addr = StringVar()
            city = StringVar()
            zipcode = StringVar()
            county = StringVar()
            country = StringVar()


            # Entry Variables
            entry1 = Entry(textvariable=cid,width=20)
            tboxes.append(entry1)
            widgets.append(entry1)
            entry1.insert(0,str(result[0][0]))
            entry1.place(x=200,y=130)

            entry2 = Entry(textvariable=name,width=20)
            tboxes.append(entry2)
            widgets.append(entry2)
            entry2.insert(0,str(result[0][1]))
            entry2.place(x=200,y=180)

            entry3 = Entry(textvariable=addr,width=20)
            tboxes.append(entry3)
            widgets.append(entry3)
            entry3.insert(0,str(result[0][2]))
            entry3.place(x=200,y=230)

            entry4 = Entry(textvariable=city,width=20)
            tboxes.append(entry4)
            widgets.append(entry4)
            entry4.insert(0,str(result[0][3]))
            entry4.place(x=200,y=280)

            entry5 = Entry(textvariable=zipcode,width=20)
            tboxes.append(entry5)
            widgets.append(entry5)
            entry5.insert(0,str(result[0][4]))
            entry5.place(x=200,y=330)

            entry6 = Entry(textvariable=county,width=20)
            tboxes.append(entry6)
            widgets.append(entry6)
            entry6.insert(0,str(result[0][5]))
            entry6.place(x=200,y=380)

            entry7 = Entry(textvariable=country,width=20)
            tboxes.append(entry7)
            widgets.append(entry7)
            entry7.insert(0,str(result[0][6]))
            entry7.place(x=200,y=430)

            save_record = Button(root,text='Update',command=update)
            widgets.append(save_record)
            tboxes.append(save_record)
            save_record.place(x=200,y=480)
            
            delete_button = Button(root,text='Delete Entry',command=delete)
            delete_button.place(x=300,y=480)
            tboxes.append(delete_button)
            widgets.append(delete_button)

         searched_label = Label(root,text=result)
         searched_label.grid(row=6,column=1,padx=10,pady=10,columnspan=2)
         widgets.append(searched_label)
         labels.append(searched_label)

         clear_button = Button(root,text='Clear Search Results',command=clear_labels)
         clear_button.grid(row=0,column=4,padx=10,pady=10)
         labels.append(clear_button)
      
      search_entry = StringVar()
      search_by = StringVar()
      
      search_box = Entry(root,textvariable=search_entry)
      search_box.grid(row=0,column=2,padx=10,pady=10)
      labels.append(search_box)

      search_type_label = Label(root,text='Search by:')
      search_type_label.grid(row=0,column=0,padx=10,pady=10)
      
      search_box_label = OptionMenu(root,search_by,*SEARCH_MENU)
      search_box_label.grid(row=0,column=1,padx=10,pady=10)
      labels.append(search_box_label)

      search_button = Button(root,text='Search Chapters',command=search_now)
      search_button.grid(row=0,column=3,padx=10,pady=10)
      labels.append(search_button)

   search_chapter()

def modify_award_details():
   destroy_all()

   SEARCH_MENU = ['Student ID']
   search_entry = StringVar()
   search_by = StringVar()

   def search_now():
      command = ''

      if search_by.get() == 'Student ID':
         command = f"select * from award_detail where student_id='{search_entry.get()}'"

      conn = cx_Oracle.connect(connstr)
      curs = conn.cursor()
      curs.execute('select student_id from student')
      SID = curs.fetchall()
      curs.execute('select community_service_id from community_service')
      CS_ID_IN = curs.fetchall()
      curs.execute(command)
      result = curs.fetchall()
      searched_label = Label(root,text='')
      print(result)

      CS_ID = list(itertools.chain(*CS_ID_IN))

      widgets = []

      # Include address twice in the sql commands
      def update():
         command2 = f"""
         update award_detail set
         has_award='{has_award.get()}' where student_id='{sid.get()}' and program_category_id='{progcat_id.get()}' and community_service_id='{cs_id.get()}'
         """
         
         print(command2)
         curs.execute(command2)
         conn.commit()

      def delete():
         command3 = f"""
         delete from award_detail where student_id='{sid.get()}' and program_category_id='{progcat_id.get()}' and community_service_id='{cs_id.get()}'
         """
         print(command3)
         curs.execute(command3)
         conn.commit()
   
      def clear():
         for widget in widgets:
            widget.destroy()
      
      if not result:
         result = 'No Results Found'

      else:
            PROGCAT_ID = ['0','1','2','3']
            HAS_AWARD = ['Y','N']
            
            label1 = Label(root, text='Program Category ID:')
            labels.append(label1)
            widgets.append(label1)
            label1.place(x=100,y=130)
            
            label2 = Label(root, text='Student ID:')
            labels.append(label2)
            widgets.append(label2)
            label2.place(x=100,y=180)
            
            label3 = Label(root,text='Community Service ID:')
            labels.append(label3)
            widgets.append(label3)
            label3.place(x=100,y=230)
            
            label4 = Label(root,text='Has Award?:')
            labels.append(label4)
            widgets.append(label4)
            label4.place(x=100,y=280)

            # Variables
            progcat_id = StringVar()
            sid = StringVar()
            cs_id = StringVar()
            has_award = StringVar()

            # Entry variables
            entry1 = Entry(root,textvariable=progcat_id)
            tboxes.append(entry1)
            widgets.append(entry1)
            entry1.insert(0,result[0][0])
            entry1.place(x=250,y=130)

            entry2 = Entry(root,textvariable=sid)
            tboxes.append(entry2)
            widgets.append(entry2)
            entry2.insert(0,str(result[0][1]))
            entry2.place(x=200,y=180)

            entry3 = Entry(root,textvariable=cs_id)
            tboxes.append(entry3)
            widgets.append(entry3)
            entry3.insert(0,result[0][2])
            entry3.place(x=250,y=230)

            entry4 = OptionMenu(root,has_award,*HAS_AWARD)
            tboxes.append(entry4)
            widgets.append(entry4)
            entry4.place(x=250,y=280)

            save_record = Button(root,text='Update',command=update)
            widgets.append(save_record)
            tboxes.append(save_record)
            save_record.place(x=200,y=480)
            
            delete_button = Button(root,text='Delete Entry',command=delete)
            delete_button.place(x=300,y=480)
            tboxes.append(delete_button)
            widgets.append(delete_button)

      searched_label = Label(root,text=result)
      searched_label.grid(row=6,column=1,padx=10,pady=10,columnspan=2)
      labels.append(searched_label)
      widgets.append(searched_label)

      clear_button = Button(root,text='Clear All Entries',command=clear)
      clear_button.grid(row=0,column=4,padx=10,pady=10)
      labels.append(clear_button)
      widgets.append(clear_button)

   
   search_by = StringVar()
   
   search_box = Entry(root,textvariable=search_entry)
   search_box.grid(row=0,column=2,padx=10,pady=10)
   labels.append(search_box)

   search_type_label = Label(root,text='Search by:')
   search_type_label.grid(row=0,column=0,padx=10,pady=10)
   labels.append(search_type_label)

   search_box_label = OptionMenu(root,search_by,*SEARCH_MENU)
   search_box_label.grid(row=0,column=1,padx=10,pady=10)
   labels.append(search_box_label)

   search_button = Button(root,text='Search Award Details',command=search_now)
   search_button.grid(row=0,column=3,padx=10,pady=10)
   labels.append(search_button)

def modify_cs():
   destroy_all()

   SEARCH_MENU = ['ID','Name']
   search_entry = StringVar()
   search_by = StringVar()

   def search_now():
      command = ''

      if search_by.get() == 'ID':
         command = f'select * from community_service where community_service_id={search_entry.get()}'
      elif search_by.get() == 'Name':
         command = f"select * from community_service where cs_name='{search_entry.get()}'"

      conn = cx_Oracle.connect(connstr)
      curs = conn.cursor()
      curs.execute(command)
      result = curs.fetchall()
      searched_label = Label(root,text='')
      print(result)

      widgets = []

      # Include address twice in the sql commands
      def update():
         command2 = f"""update community_service set cs_name='{name.get()}',cs_addr='{addr.get()}',cs_contact_person_full_name='{contact.get()}',
         address='{addr.get()}',city='{city.get()}',zipcode='{zipcode.get()}',county='{county.get()}',country='{country.get()}',work_phone='{wphone.get()}',
         cell_phone='{cphone.get()}',email='{email.get()}' where community_service_id='{csid.get()}'
         """
         print(command2)
         curs.execute(command2)
         conn.commit()

      def delete():
         command3 = f"""
         delete from community_service where community_service_id={csid.get()}
         """

      def clear():
         for widget in widgets:
            widget.destroy()
      
      if not result:
         result = 'No Results Found'

      else:
            widgets = []
         
            label1 = Label(root, text='ID:')
            labels.append(label1)
            widgets.append(label1)
            label1.place(x=100,y=130)
            
            label2 = Label(root, text='Name:')
            labels.append(label2)
            widgets.append(label2)
            label2.place(x=100,y=180)
            
            label3 = Label(root,text='Address:')
            labels.append(label3)
            widgets.append(label3)
            label3.place(x=100,y=230)
            
            label4 = Label(root,text='Contact (Full Name):')
            labels.append(label4)
            widgets.append(label4)
            label4.place(x=100,y=280)
            
            label5 = Label(root,text='City:')
            labels.append(label5)
            widgets.append(label5)
            label5.place(x=100,y=330)

            label6 = Label(root,text='Zipcode:')
            labels.append(label6)
            widgets.append(label6)
            label6.place(x=100,y=380)

            label7 = Label(root,text='County:')
            labels.append(label7)
            widgets.append(label7)
            label7.place(x=100,y=430)

            label8 = Label(root,text='Country:')
            labels.append(label8)
            widgets.append(label8)
            label8.place(x=100,y=480)

            label9 = Label(root,text='Work Phone:')
            labels.append(label9)
            widgets.append(label9)
            label9.place(x=100,y=530)

            label10 = Label(root,text='Cell Phone:')
            labels.append(label10)
            widgets.append(label10)
            label10.place(x=100,y=580)

            label11 = Label(root,text='Email:')
            labels.append(label11)
            widgets.append(label11)
            label11.place(x=100,y=630)

            # Variables
            csid = StringVar()
            name = StringVar()
            addr = StringVar()
            contact = StringVar()
            city = StringVar()
            zipcode = StringVar()
            county = StringVar()
            country = StringVar()
            wphone = StringVar()
            cphone = StringVar()
            email = StringVar()

            # Entry variables
            entry1 = Entry(textvariable=csid,width=20)
            tboxes.append(entry1)
            widgets.append(entry1)
            entry1.insert(0,str(result[0][0]))
            entry1.place(x=200,y=130)

            entry2 = Entry(textvariable=name,width=20)
            tboxes.append(entry2)
            widgets.append(entry2)
            entry2.insert(0,str(result[0][1]))
            entry2.place(x=200,y=180)

            entry3 = Entry(textvariable=addr,width=20)
            tboxes.append(entry3)
            widgets.append(entry3)
            entry3.insert(0,str(result[0][2]))
            entry3.place(x=200,y=230)

            entry4 = Entry(textvariable=contact,width=20)
            tboxes.append(entry4)
            widgets.append(entry4)
            entry4.insert(0,str(result[0][3]))
            entry4.place(x=200,y=280)

            entry5 = Entry(textvariable=city,width=20)
            tboxes.append(entry5)
            widgets.append(entry5)
            entry5.insert(0,str(result[0][4]))
            entry5.place(x=200,y=330)

            entry6 = Entry(textvariable=zipcode,width=20)
            tboxes.append(entry6)
            widgets.append(entry6)
            entry6.insert(0,str(result[0][5]))
            entry6.place(x=200,y=380)

            entry7 = Entry(textvariable=county,width=20)
            tboxes.append(entry7)
            widgets.append(entry7)
            entry7.insert(0,str(result[0][6]))
            entry7.place(x=200,y=430)

            entry8 = Entry(textvariable=country,width=20)
            tboxes.append(entry8)
            widgets.append(entry8)
            entry8.insert(0,str(result[0][5]))
            entry8.place(x=200,y=480)

            entry8 = Entry(textvariable=wphone,width=20)
            tboxes.append(entry8)
            widgets.append(entry8)
            entry8.insert(0,str(result[0][6]))
            entry8.place(x=200,y=530)

            entry9 = Entry(textvariable=cphone,width=20)
            tboxes.append(entry9)
            widgets.append(entry9)
            entry9.insert(0,str(result[0][5]))
            entry9.place(x=200,y=580)

            entry10 = Entry(textvariable=email,width=20)
            tboxes.append(entry10)
            widgets.append(entry10)
            entry10.insert(0,str(result[0][6]))
            entry10.place(x=200,y=630)

            save_record = Button(root,text='Update',command=update)
            widgets.append(save_record)
            tboxes.append(save_record)
            save_record.place(x=200,y=680)
            
            delete_button = Button(root,text='Delete Entry',command=delete)
            delete_button.place(x=300,y=680)
            tboxes.append(delete_button)
            widgets.append(delete_button)

      searched_label = Label(root,text=result)
      searched_label.grid(row=6,column=1,padx=10,pady=10,columnspan=2)
      labels.append(searched_label)
      widgets.append(searched_label)

      clear_button = Button(root,text='Clear All Entries',command=clear)
      clear_button.grid(row=0,column=4,padx=10,pady=10)
      labels.append(clear_button)
      widgets.append(clear_button)
   
   search_by = StringVar()
   
   search_box = Entry(root,textvariable=search_entry)
   search_box.grid(row=0,column=2,padx=10,pady=10)
   labels.append(search_box)

   search_type_label = Label(root,text='Search by:')
   search_type_label.grid(row=0,column=0,padx=10,pady=10)
   labels.append(search_type_label)

   search_box_label = OptionMenu(root,search_by,*SEARCH_MENU)
   search_box_label.grid(row=0,column=1,padx=10,pady=10)
   labels.append(search_box_label)

   search_button = Button(root,text='Search Entries',command=search_now)
   search_button.grid(row=0,column=3,padx=10,pady=10)
   labels.append(search_button)

def modify_cshrs():
   destroy_all()

   SEARCH_MENU = ['Student ID','Program Category ID']
   search_entry = StringVar()
   search_by = StringVar()

   widgets = []

   def search_now():
      command = ''
      
      if search_by.get() == 'Student ID':
         command = f'select * from csa_hour where student_id = {search_entry.get()}'   
      elif search_by.get() == 'Program Category ID':
         command = f'select * from csa_hour where program_category_id = {search_entry.get()}'
         
      conn = cx_Oracle.connect(connstr)
      curs = conn.cursor()
      curs.execute(command)
      result = curs.fetchall()
      searched_label = Label(root,text='')
      print(result)

      def update():
         command2 = f"""
         update csa_hour set csa_hour_id='{cshr_id.get()}',student_id='{sid.get()}',program_category_id='{progcat_id.get()}',
         community_service_id='{cs_id.get()}',hours_worked='{hrs_worked.get()}',
         date_worked=TO_DATE('{date.get()}', 'YYYY-MM-DD HH24:MI:SS') where csa_hour_id='{cshr_id.get()}'
         """
         print(command2)
         curs.execute(command2)
         conn.commit()

      def delete():
         command3 = f"""
         delete from csa_hour where csa_hour = 'cshr_id.get()'
         """
      
      def clear():
         for widget in widgets:
            widget.destroy()

      if not result:
         result = 'No Results Found'

      else:
            # Label variables
            label1 = Label(root, text='Community Service Hour ID:')
            labels.append(label1)
            widgets.append(label1)
            label1.place(x=100,y=130)
            
            label2 = Label(root, text='Student ID:')
            labels.append(label2)
            widgets.append(label2)
            label2.place(x=100,y=180)
            
            label3 = Label(root,text='Program Category ID:')
            labels.append(label3)
            widgets.append(label3)
            label3.place(x=100,y=230)
            
            label4 = Label(root,text='Community Service ID:')
            labels.append(label4)
            widgets.append(label4)
            label4.place(x=100,y=280)

            label5 = Label(root,text='Hours Worked:')
            labels.append(label5)
            widgets.append(label5)
            label5.place(x=100,y=330)

            label6 = Label(root,text='Date:')
            labels.append(label6)
            widgets.append(label6)
            label6.place(x=100,y=380)

            # Variables
            cshr_id = StringVar()
            sid = StringVar()
            progcat_id = StringVar()
            cs_id = StringVar()
            hrs_worked = StringVar()
            date = StringVar()

            # Entry Variables
            entry1 = Entry(textvariable=cshr_id,width=20)
            tboxes.append(entry1)
            widgets.append(entry1)
            entry1.insert(0,str(result[0][0]))
            entry1.place(x=270,y=130)

            entry2 = Entry(textvariable=sid,width=20)
            tboxes.append(entry2)
            widgets.append(entry2)
            entry2.insert(0,str(result[0][1]))
            entry2.place(x=270,y=180)

            entry3 = Entry(textvariable=progcat_id,width=20)
            tboxes.append(entry3)
            widgets.append(entry3)
            entry3.insert(0,str(result[0][2]))
            entry3.place(x=270,y=230)

            entry4 = Entry(textvariable=cs_id,width=20)
            tboxes.append(entry4)
            widgets.append(entry4)
            entry4.insert(0,str(result[0][3]))
            entry4.place(x=270,y=280)

            entry5 = Entry(textvariable=hrs_worked,width=20)
            tboxes.append(entry5)
            widgets.append(entry5)
            entry5.insert(0,str(result[0][4]))
            entry5.place(x=270,y=330)

            entry6 = Entry(textvariable=date,width=20)
            tboxes.append(entry6)
            widgets.append(entry6)
            entry6.insert(0,str(result[0][5]))
            entry6.place(x=270,y=380)

            save_record = Button(root,text='Update',command=update)
            widgets.append(save_record)
            tboxes.append(save_record)
            save_record.place(x=200,y=500)
            
            delete_button = Button(root,text='Delete Entry',command=delete)
            delete_button.place(x=300,y=500)
            tboxes.append(delete_button)
            widgets.append(delete_button)

      searched_label = Label(root,text=result)
      searched_label.grid(row=6,column=1,padx=10,pady=10,columnspan=2)
      labels.append(searched_label)
      widgets.append(searched_label)

      clear_button = Button(root,text='Clear All Entries',command=clear)
      clear_button.grid(row=0,column=4,padx=10,pady=10)
      labels.append(clear_button)
      widgets.append(clear_button)
      
   search_by = StringVar()
   
   search_box = Entry(root,textvariable=search_entry)
   search_box.grid(row=0,column=2,padx=10,pady=10)
   labels.append(search_box)

   search_type_label = Label(root,text='Search by:')
   search_type_label.grid(row=0,column=0,padx=10,pady=10)
   labels.append(search_type_label)

   search_box_label = OptionMenu(root,search_by,*SEARCH_MENU)
   search_box_label.grid(row=0,column=1,padx=10,pady=10)
   labels.append(search_box_label)

   search_button = Button(root,text='Search Entries',command=search_now)
   search_button.grid(row=0,column=3,padx=10,pady=10)
   labels.append(search_button)

def print_report():
   destroy_all()

   def create_report():
      conn = cx_Oracle.connect(connstr)
      curs = conn.cursor()
      command = f"""
      select  st.first_name, st.last_name, 
      pc.name Program_name,
      cs.cs_name,
      ch.HOURS_WORKED,
      ch.DATE_WORKED
      from csa_hour ch, 
        student st,
        program_category pc,
        community_service cs
      where st.student_id = ch.student_id
      and pc.program_category_id = ch.program_category_id
      and cs.community_service_id = ch.community_service_id
      and date_worked between '{d1.get()}' and '{d2.get()}'
      """
      print(command)
      curs.execute(command)
      #print(f"select * from csa_hour where date_worked between '{d1.get()}' and '{d2.get()}'")
      result1 = curs.fetchall()
      result2 = list(itertools.chain(result1))
      
      document = Document()

      document.add_heading('FBLA Community Service Awards Report',0)

      section = document.sections[-1]
      section.orientation = WD_ORIENT.LANDSCAPE
      section.page_width = 10058400
      section.page_height = 7772400
      
      table = document.add_table(rows=1,cols=6)
      hdr_cells = table.rows[0].cells
      hdr_cells[0].text = 'First Name'
      hdr_cells[1].text = 'Last Name'
      hdr_cells[2].text = 'Program Category'
      hdr_cells[3].text = 'CS Name'
      hdr_cells[4].text = 'Hours Worked'
      hdr_cells[5].text = 'Date Worked'
      
      for cid, sid, pcid, csid, hrs, date in result2:
         row_cells = table.add_row().cells
         row_cells[0].text = str(cid)
         row_cells[1].text = str(sid)
         row_cells[2].text = str(pcid)
         row_cells[3].text = str(csid)
         row_cells[4].text = str(hrs)
         row_cells[5].text = str(date)

      document.add_page_break()

      document.save('fbla-csa-report-2020.docx')

   label1 = Label(root,text='Please submit two dates to query the data added to the report')
   labels.append(label1)
   label1.grid(row=0,column=0)

   label2 = Label(root,text='Start Date (DD-MON-YYYY):')
   labels.append(label2)
   label2.grid(row=1,column=0)

   label3 = Label(root,text='End Date (DD-MON-YYYY):')
   labels.append(label3)
   label3.grid(row=2,column=0)

   d1 = StringVar()
   d2 = StringVar()

   entry1 = Entry(root,textvariable=d1)
   labels.append(entry1)
   entry1.grid(row=1,column=1)
   
   entry2 = Entry(root,textvariable=d2)
   labels.append(entry2)
   entry2.grid(row=2,column=1)
   
   button = Button(root,text='Create Report',command=create_report)
   labels.append(button)
   button.grid(row=4,column=1)

def backup():
   destroy_all()
   
   def backup_files():
      os.system('D:\Python_Files\FBLA_MAIN\fbla_csa.bat')
      
      label = Label(root,text="All data is being backed up in 'D:\Python_Files\FBLA_MAIN\FBLA_CSA.DMP'")
      labels.append(label)
      label.grid(row=1,column=2)


   button = Button(root,text='Backup All FBLA Data',command=backup_files)
   labels.append(button)
   button.grid(row=0,column=2)

root = Tk()
root.title('FBLA Community Service Hours Program')
width = root.winfo_screenwidth()
height = root.winfo_screenheight()
root.geometry("%dx%d" % (width, height))

frame2 = Frame(root)
frame2.grid(row=0,column=0)

menubar = Menu(root)
filemenu = Menu(menubar, tearoff=0)
filemenu.add_command(label="Student", command=add_student)
filemenu.add_command(label="Chapter", command=add_chapter)
filemenu.add_command(label="Award Detail", command=add_award)
filemenu.add_command(label="Community Service", command=add_cs)
filemenu.add_command(label="Coummunity Service Hours", command=add_cshrs)
#filemenu.add_command(label="Program Category", command=donothing)
#filemenu.add_command(label="Teacher (Administrator)", command=donothing)

menubar.add_cascade(label="Add Data", menu=filemenu)
editmenu = Menu(menubar, tearoff=0)
editmenu.add_command(label="Student", command=modify_student)
editmenu.add_command(label="Chapter", command=modify_chapter)
editmenu.add_command(label="Award Detail", command=modify_award_details)
editmenu.add_command(label="Community Service", command=modify_cs)
editmenu.add_command(label="Coummunity Service Hours", command=modify_cshrs)
#editmenu.add_command(label="Program Category", command=donothing)
menubar.add_cascade(label="Modify Data", menu=editmenu)
#editmenu.add_command(label="Teacher (Administrator)", command=donothing)
#editmenu.add_separator()
#editmenu.add_command(label="Cut", command=donothing)
#editmenu.add_command(label="Copy", command=donothing)
#editmenu.add_command(label="Paste", command=donothing)
#editmenu.add_command(label="Delete", command=donothing)
#editmenu.add_command(label="Select All", command=donothing)

helpmenu = Menu(menubar, tearoff=0)
helpmenu.add_command(label="Chapter", command=query_chapter)
helpmenu.add_command(label="Student", command=query_student)
helpmenu.add_command(label="Program Category", command=query_progcat)
helpmenu.add_command(label="Community Service", command=query_csa)
helpmenu.add_command(label="Coummunity Service Hours", command=query_csa_hour)
helpmenu.add_command(label="Award Detail", command=query_award_det)
#helpmenu.add_command(label="Teacher (Adminstrator)", command=donothing)
menubar.add_cascade(label="Search Data", menu=helpmenu)

printmenu = Menu(menubar, tearoff=0)
printmenu.add_command(label='Create Report',command=print_report)
#printmenu.add_command(label='Print Monthly Report',command=month_report)
menubar.add_cascade(label='Print Report',menu=printmenu)

exitmenu = Menu(menubar, tearoff=0)
menubar.add_cascade(label="Options", menu=exitmenu)
exitmenu.add_command(label='Backup FBLA Data',command=backup)
exitmenu.add_command(label="Quit", command=root.destroy)

root.config(menu=menubar)
root.mainloop()
